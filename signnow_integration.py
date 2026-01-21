"""
SignNow API Integration Module
Handles document creation and sending for contract signing
"""

import requests
import json
import base64
from typing import Dict, Optional, Tuple
import streamlit as st
import os
from docx_template_processor import DocxTemplateProcessor

class SignNowAPI:
    """SignNow API client for document creation and sending"""
    
    def __init__(self, client_id: str, client_secret: str, basic_auth_token: str, 
                 username: str, password: str, api_key: str = None):
        """
        Initialize SignNow API client
        
        Args:
            client_id: SignNow application client ID
            client_secret: SignNow application client secret
            basic_auth_token: SignNow basic authorization token
            username: SignNow account username/email
            password: SignNow account password
            api_key: SignNow API key (optional)
        """
        self.client_id = client_id
        self.client_secret = client_secret
        self.basic_auth_token = basic_auth_token
        self.username = username
        self.password = password
        self.api_key = api_key
        self.base_url = "https://api.signnow.com"  # Updated to production URL
        self.access_token = None
        self.user_email = None
        self.docx_processor = DocxTemplateProcessor()
        
    def authenticate(self) -> bool:
        """
        Authenticate with SignNow API using OAuth2 password grant
        
        Returns:
            bool: True if authentication successful, False otherwise
        """
        try:
            # OAuth2 token endpoint
            auth_url = "https://api.signnow.com/oauth2/token"
            
            # Headers as per SignNow documentation
            headers = {
                "Authorization": f"Basic {self.basic_auth_token}",
                "Content-Type": "application/x-www-form-urlencoded"
            }
            
            # Request body as per SignNow documentation
            auth_data = {
                "username": self.username,
                "password": self.password,
                "grant_type": "password",
                "scope": "*"
            }
            
            response = requests.post(auth_url, headers=headers, data=auth_data, timeout=30)
            response.raise_for_status()
            
            auth_response = response.json()
            self.access_token = auth_response.get("access_token")
            
            if not self.access_token:
                st.error("Failed to get access token from SignNow")
                return False
            
            # Get user email for document sending
            self._get_user_email()
                
            return True
            
        except requests.exceptions.RequestException as e:
            st.error(f"SignNow authentication failed: {str(e)}")
            return False
        except Exception as e:
            st.error(f"Unexpected error during SignNow authentication: {str(e)}")
            return False
    
    def _get_user_email(self) -> bool:
        """
        Get user's email address for document sending
        
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            user_url = f"{self.base_url}/user"
            
            headers = {
                "Authorization": f"Bearer {self.access_token}",
                "Content-Type": "application/json"
            }
            
            response = requests.get(user_url, headers=headers, timeout=30)
            response.raise_for_status()
            
            user_data = response.json()
            emails = user_data.get('emails', [])
            
            if emails:
                self.user_email = emails[0]
                return True
            else:
                st.error("No email found for user")
                return False
                
        except Exception as e:
            st.error(f"Failed to get user email: {str(e)}")
            return False
    
    def create_document_from_template(self, template_type: str, document_name: str, 
                                   client_name: str, email: str, 
                                   contract_amount: str = None, contract_date: str = None) -> Optional[str]:
        """
        Create a new document by processing the original .docx template with exact formatting preserved
        
        Args:
            template_type: Type of template to use (development_contract, development_terms, terms_conditions, production_contract)
            document_name: Name for the new document
            client_name: Client name value
            email: Email value
            contract_amount: Contract amount value (optional)
            contract_date: Contract date value (optional)
            
        Returns:
            str: Document ID if successful, None otherwise
        """
        if not self.access_token:
            if not self.authenticate():
                return None
        
        try:
            # Process the .docx template with exact formatting preserved
            processed_docx_path = self.docx_processor.process_document(
                template_type=template_type,
                client_name=client_name,
                email=email,
                contract_amount=contract_amount,
                contract_date=contract_date
            )
            
            # Convert .docx to PDF for SignNow upload
            pdf_content = self._convert_docx_to_pdf(processed_docx_path)
            
            # Upload the PDF to SignNow
            files = {
                'file': (f'{document_name}.pdf', pdf_content, 'application/pdf')
            }
            
            data = {
                'name': document_name
            }
            
            headers = {
                "Authorization": f"Bearer {self.access_token}"
            }
            
            response = requests.post(
                f"{self.base_url}/document",
                files=files,
                data=data,
                headers=headers,
                timeout=60
            )
            response.raise_for_status()
            
            create_response = response.json()
            document_id = create_response.get("id")
            
            if not document_id:
                st.error("Failed to create document")
                return None
                
            # Clean up temporary file
            if os.path.exists(processed_docx_path):
                os.remove(processed_docx_path)
                
            return document_id
            
        except requests.exceptions.RequestException as e:
            st.error(f"Failed to create document: {str(e)}")
            return None
        except Exception as e:
            st.error(f"Unexpected error creating document: {str(e)}")
            return None
    
    def _convert_docx_to_pdf(self, docx_path: str, highlight_values: list | None = None) -> bytes:
        """
        Convert .docx file to PDF bytes with proper formatting preservation using ReportLab
        
        Args:
            docx_path: Path to the .docx file
            
        Returns:
            bytes: PDF content
        """
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from io import BytesIO
            import html as html_escape
            
            doc = Document(docx_path)
            
            # Create PDF buffer
            buffer = BytesIO()
            pdf_doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch, bottomMargin=1*inch)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles for better formatting
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                spaceAfter=20,
                alignment=1,  # Center alignment
                textColor=colors.black,
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                fontName='Helvetica-Bold',
                textColor=colors.black
            )
            
            bold_style = ParagraphStyle(
                'CustomBold',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=8,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                fontName='Helvetica'
            )
            
            amount_style = ParagraphStyle(
                'ContractAmount',
                parent=styles['Normal'],
                fontSize=14,
                spaceAfter=8,
                fontName='Helvetica-Bold',
                textColor=colors.green
            )
            
            placeholder_style = ParagraphStyle(
                'Placeholder',
                parent=styles['Normal'],
                fontSize=16,
                spaceAfter=12,
                alignment=1,  # Center alignment
                fontName='Helvetica-Bold',
                textColor=colors.blue
            )
            
            screenshot_placeholder_style = ParagraphStyle(
                'ScreenshotPlaceholder',
                parent=styles['Normal'],
                fontSize=16,
                spaceAfter=12,
                alignment=1,  # Center alignment
                fontName='Helvetica-Bold',
                textColor=colors.red
            )
            
            # Normalize highlight values
            highlight_values = [v for v in (highlight_values or []) if isinstance(v, str) and v.strip()]

            # Build PDF content
            story = []
            
            # Check for images/logos first
            if hasattr(doc, 'inline_shapes') and doc.inline_shapes:
                story.append(Paragraph("IMAGE/LOGO NOT DISPLAYED IN PREVIEW", placeholder_style))
                story.append(Spacer(1, 20))
            
            # Add document content with exact formatting preservation
            for paragraph in doc.paragraphs:
                # Check if this paragraph contains an image (screenshot placeholder)
                has_image = False
                if paragraph.runs:
                    for run in paragraph.runs:
                        if run._element.xpath('.//a:blip'):
                            has_image = True
                            break
                
                # If paragraph has an image, show red placeholder
                if has_image:
                    story.append(Paragraph("IMAGE ATTACHED GOES HERE", screenshot_placeholder_style))
                    story.append(Spacer(1, 12))
                    continue
                
                if paragraph.text.strip():
                    text = paragraph.text.strip()
                    
                    # Check for actual bold formatting in the document
                    is_bold = paragraph.runs and any(run.bold for run in paragraph.runs)
                    
                    # Check for actual font size in the document
                    font_size = 11  # Default size
                    if paragraph.runs:
                        for run in paragraph.runs:
                            if run.font.size:
                                font_size = run.font.size.pt
                                break
                    
                    # Map font families to ReportLab supported fonts
                    def map_font_family(font_name, is_bold=False):
                        """Map document fonts to ReportLab supported fonts"""
                        font_name = font_name.lower() if font_name else 'helvetica'
                        
                        # Font mapping
                        font_map = {
                            'arial': 'Helvetica',
                            'times': 'Times-Roman',
                            'times new roman': 'Times-Roman',
                            'courier': 'Courier',
                            'helvetica': 'Helvetica',
                            'calibri': 'Helvetica',  # Map Calibri to Helvetica
                            'verdana': 'Helvetica',  # Map Verdana to Helvetica
                        }
                        
                        # Get base font
                        base_font = font_map.get(font_name, 'Helvetica')
                        
                        # Add bold suffix if needed
                        if is_bold:
                            if base_font == 'Helvetica':
                                return 'Helvetica-Bold'
                            elif base_font == 'Times-Roman':
                                return 'Times-Bold'
                            elif base_font == 'Courier':
                                return 'Courier-Bold'
                            else:
                                return 'Helvetica-Bold'  # Fallback
                        else:
                            return base_font
                    
                    # Check for actual font family in the document
                    font_family = 'Helvetica'  # Default font
                    if paragraph.runs:
                        for run in paragraph.runs:
                            if run.font.name:
                                font_family = run.font.name
                                break
                    
                    # Force all paragraphs to normal weight; only bold exact variable values inline
                    style = ParagraphStyle(
                        'DynamicNormal',
                        parent=styles['Normal'],
                        fontSize=font_size,
                        spaceAfter=4,
                        fontName=map_font_family(font_family, False),
                        textColor=colors.black
                    )
                    
                    # Escape text for XML/HTML and bold only highlight values inline
                    escaped_text = html_escape.escape(text)
                    if highlight_values:
                        # Replace longer values first to avoid partial nesting
                        for val in sorted(highlight_values, key=lambda v: len(v or ''), reverse=True):
                            if not val:
                                continue
                            escaped_val = html_escape.escape(val)
                            escaped_text = escaped_text.replace(escaped_val, f"<b>{escaped_val}</b>")

                    story.append(Paragraph(escaped_text, style))
                    
                    # Add spacing based on paragraph formatting
                    if paragraph.paragraph_format.space_after:
                        spacing = paragraph.paragraph_format.space_after.pt
                        story.append(Spacer(1, spacing))
                    else:
                        story.append(Spacer(1, 2))  # Minimal default spacing
            
            # Add table placeholders
            if doc.tables:
                story.append(Spacer(1, 20))
                story.append(Paragraph("TABLE NOT DISPLAYED IN PREVIEW", placeholder_style))
            
            # Build PDF
            pdf_doc.build(story)
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            return pdf_bytes
            
        except Exception as e:
            st.error(f"Error converting .docx to PDF: {str(e)}")
            # Fallback to simple PDF
            return b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n%%EOF"

    def _merge_pdfs(self, pdf_bytes_list: list[bytes]) -> bytes:
        """Merge multiple PDF byte streams into a single PDF bytes using PyPDF2."""
        try:
            from PyPDF2 import PdfReader, PdfWriter
            from io import BytesIO

            writer = PdfWriter()
            for pdf_bytes in pdf_bytes_list:
                reader = PdfReader(BytesIO(pdf_bytes))
                for page in reader.pages:
                    writer.add_page(page)

            out = BytesIO()
            writer.write(out)
            out.seek(0)
            return out.getvalue()
        except Exception as e:
            st.error(f"Failed to merge PDFs: {str(e)}")
            # Return first PDF as fallback
            return pdf_bytes_list[0] if pdf_bytes_list else b""

    def upload_pdf(self, document_name: str, pdf_bytes: bytes) -> str | None:
        """Upload a single PDF to SignNow and return document id."""
        if not self.access_token:
            if not self.authenticate():
                return None
        try:
            files = {
                'file': (f'{document_name}.pdf', pdf_bytes, 'application/pdf')
            }
            data = { 'name': document_name }
            headers = { "Authorization": f"Bearer {self.access_token}" }
            response = requests.post(f"{self.base_url}/document", files=files, data=data, headers=headers, timeout=90)
            response.raise_for_status()
            doc_id = response.json().get("id")
            return doc_id
        except Exception as e:
            st.error(f"Failed to upload merged PDF: {str(e)}")
            return None

    def create_and_send_merged_pair(self, pair_type: str, contract_docx_path: str, terms_docx_path: str,
                                    document_name: str, email: str, highlight_values: list | None = None) -> tuple[bool, str]:
        """Convert both DOCX to PDFs, merge into one, upload, and send invite."""
        try:
            # Convert to PDFs
            contract_pdf = self._convert_docx_to_pdf(contract_docx_path, highlight_values=highlight_values)
            terms_pdf = self._convert_docx_to_pdf(terms_docx_path, highlight_values=highlight_values)
            # Merge PDFs: contract first, then terms
            merged_pdf = self._merge_pdfs([contract_pdf, terms_pdf])
            if not merged_pdf:
                return False, "Failed to generate merged PDF"
            # Upload merged PDF with provided document_name (already formatted by caller)
            doc_id = self.upload_pdf(document_name, merged_pdf)
            if not doc_id:
                return False, "Failed to upload merged document"
            # Send for signing
            if self.send_document_for_signing(doc_id, email, document_name, ""):
                return True, f"Merged document sent successfully to {email}. Document ID: {doc_id}"
            return False, "Merged document uploaded but failed to send for signing"
        except Exception as e:
            return False, f"Error creating/sending merged pair: {str(e)}"

    def _trigger_field_extraction(self, document_id: str) -> bool:
        """
        Trigger SignNow to extract fields from text tags if needed.
        After uploading a DOCX with text tags, SignNow may need to process them.
        
        Args:
            document_id: ID of the document
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Some SignNow implementations require field extraction endpoint
            # Try: POST /document/{document_id}/fieldextract
            extract_url = f"{self.base_url}/document/{document_id}/fieldextract"
            
            headers = {
                "Authorization": f"Bearer {self.access_token}",
                "Content-Type": "application/json"
            }
            
            # Trigger field extraction
            response = requests.post(extract_url, headers=headers, timeout=60)
            
            # 200 or 202 are acceptable (202 = accepted/processing)
            if response.status_code in [200, 202]:
                return True
            else:
                # Field extraction endpoint might not exist or work differently
                return False
                
        except requests.exceptions.RequestException as e:
            # Endpoint might not exist - that's okay, SignNow may process text tags automatically
            return False
        except Exception as e:
            return False
    
    def _convert_simple_text_tags_to_signnow_format(self, doc):
        """
        Convert simple text tags to SignNow's required format.
        
        Converts:
        - {{s_Signature_Signer1}} -> {{t:s;r:y;o:"Signer 1";}}
        - {{s_Signature_Signer2}} -> {{t:s;r:y;o:"Signer 2";}}
        
        SignNow requires the complex format for automatic field creation.
        
        Args:
            doc: python-docx Document object
        """
        import re
        
        replacements_made = []
        
        # SignNow required format with size parameters (w:width, h:height)
        signer1_tag = '{{t:s;r:y;o:"Signer 1";w:200;h:25;}}'
        signer2_tag = '{{t:s;r:y;o:"Signer 2";w:200;h:25;}}'
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            original_text = paragraph.text
            new_text = original_text
            
            # Replace simple format with SignNow format
            if '{{s_Signature_Signer1}}' in new_text:
                new_text = new_text.replace('{{s_Signature_Signer1}}', signer1_tag)
                replacements_made.append(f"Para {para_idx}: Converted {{s_Signature_Signer1}} to SignNow format")
            
            if '{{s_Signature_Signer2}}' in new_text:
                new_text = new_text.replace('{{s_Signature_Signer2}}', signer2_tag)
                replacements_made.append(f"Para {para_idx}: Converted {{s_Signature_Signer2}} to SignNow format")
            
            # Apply replacement if text changed
            if new_text != original_text:
                # Preserve paragraph formatting
                paragraph.clear()
                if paragraph.runs:
                    # Try to preserve formatting from first run
                    first_run = paragraph.runs[0]
                    run = paragraph.add_run(new_text)
                    if first_run.bold:
                        run.bold = True
                    if first_run.italic:
                        run.italic = True
                    if first_run.font.size:
                        run.font.size = first_run.font.size
                    if first_run.font.name:
                        run.font.name = first_run.font.name
                else:
                    paragraph.add_run(new_text)
        
        # Also check in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        original_text = paragraph.text
                        new_text = original_text
                        
                        if '{{s_Signature_Signer1}}' in new_text:
                            new_text = new_text.replace('{{s_Signature_Signer1}}', signer1_tag)
                            replacements_made.append(f"Table {table_idx}, Cell ({row_idx},{col_idx}), Para {para_idx}: Converted {{s_Signature_Signer1}}")
                        
                        if '{{s_Signature_Signer2}}' in new_text:
                            new_text = new_text.replace('{{s_Signature_Signer2}}', signer2_tag)
                            replacements_made.append(f"Table {table_idx}, Cell ({row_idx},{col_idx}), Para {para_idx}: Converted {{s_Signature_Signer2}}")
                        
                        if new_text != original_text:
                            paragraph.clear()
                            paragraph.add_run(new_text)
        
        # Conversion happens silently - no debug output needed
        # Text tags are automatically converted to SignNow format with size parameters
    
    def _replace_placeholders_with_text_tags(self, doc):
        """
        Replace CLIENT_SIGNATURE and SALESMAN_SIGNATURE placeholders with SignNow Text Tags.
        SignNow will automatically convert these tags to signature fields at exact positions.
        
        Text Tag Format (SignNow Simple Text Tags):
        - {{s_Signature_Signer1}} = Required signature field for Signer 1
        - {{s_Signature_Signer2}} = Required signature field for Signer 2
        
        Or alternative format:
        - {{t:s;r:y;o:"Signer 1";}} = Required signature field for Signer 1
        - {{t:s;r:y;o:"Signer 2";}} = Required signature field for Signer 2
        
        Args:
            doc: python-docx Document object
        """
        import re
        
        # Use SignNow Simple Text Tags format
        # Format: {{s_Signature_Signer1}} or {{s_Signature_Signer2}}
        client_tag = '{{s_Signature_Signer1}}'
        salesman_tag = '{{s_Signature_Signer2}}'
        
        replacements_made = []
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            original_text = paragraph.text
            new_text = original_text
            
            # Replace CLIENT_SIGNATURE with SignNow text tag
            if 'CLIENT_SIGNATURE' in original_text.upper():
                new_text = re.sub(r'CLIENT_SIGNATURE', client_tag, new_text, flags=re.IGNORECASE)
                replacements_made.append(f"Para {para_idx}: CLIENT_SIGNATURE -> {client_tag}")
            
            # Replace SALESMAN_SIGNATURE with SignNow text tag
            if 'SALESMAN_SIGNATURE' in original_text.upper():
                new_text = re.sub(r'SALESMAN_SIGNATURE', salesman_tag, new_text, flags=re.IGNORECASE)
                replacements_made.append(f"Para {para_idx}: SALESMAN_SIGNATURE -> {salesman_tag}")
            
            # Apply replacement if text changed
            if new_text != original_text:
                # Preserve paragraph formatting by replacing runs
                paragraph.clear()
                # Preserve formatting from first run if available
                if paragraph.runs:
                    first_run_format = paragraph.runs[0]
                    run = paragraph.add_run(new_text)
                    if first_run_format.bold:
                        run.bold = True
                    if first_run_format.italic:
                        run.italic = True
                    if first_run_format.font.size:
                        run.font.size = first_run_format.font.size
                else:
                    paragraph.add_run(new_text)
        
        # Also check in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        original_text = paragraph.text
                        new_text = original_text
                        
                        if 'CLIENT_SIGNATURE' in original_text.upper():
                            new_text = re.sub(r'CLIENT_SIGNATURE', client_tag, new_text, flags=re.IGNORECASE)
                            replacements_made.append(f"Table {table_idx}, Cell ({row_idx},{col_idx}), Para {para_idx}: CLIENT_SIGNATURE -> {client_tag}")
                        
                        if 'SALESMAN_SIGNATURE' in original_text.upper():
                            new_text = re.sub(r'SALESMAN_SIGNATURE', salesman_tag, new_text, flags=re.IGNORECASE)
                            replacements_made.append(f"Table {table_idx}, Cell ({row_idx},{col_idx}), Para {para_idx}: SALESMAN_SIGNATURE -> {salesman_tag}")
                        
                        if new_text != original_text:
                            paragraph.clear()
                            paragraph.add_run(new_text)
        
        if replacements_made:
            st.write(f"✅ Replaced {len(replacements_made)} placeholder(s) with SignNow Text Tags")
            for replacement in replacements_made[:5]:  # Show first 5
                st.write(f"  - {replacement}")
            if len(replacements_made) > 5:
                st.write(f"  ... and {len(replacements_made) - 5} more")
        else:
            st.warning("⚠️ No signature placeholders found to replace with Text Tags")
    
    def _merge_docx(self, first_path: str, second_path: str) -> str:
        """Merge two DOCX files into one DOCX (simple append of body content). Returns merged path."""
        from docx import Document as Docx
        import os
        merged = Docx(first_path)
        second = Docx(second_path)

        # Append paragraphs and tables from second into merged
        for element in second.element.body:
            merged.element.body.append(element)

        out_dir = os.path.join(os.getcwd(), 'processed_documents')
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, 'merged_contract_terms.docx')
        
        # Text Tags are already in the documents ({{s_Signature_Signer1}} and {{s_Signature_Signer2}})
        # Convert simple text tags to SignNow's required format
        # {{s_Signature_Signer1}} -> {{t:s;r:y;o:"Signer 1";}}
        # {{s_Signature_Signer2}} -> {{t:s;r:y;o:"Signer 2";}}
        self._convert_simple_text_tags_to_signnow_format(merged)
        
        merged.save(out_path)
        return out_path

    def upload_docx(self, document_name: str, docx_path: str, use_field_extraction: bool = True) -> str | None:
        """
        Upload a DOCX to SignNow with text tag field extraction.
        
        Args:
            document_name: Name for the document
            docx_path: Path to the DOCX file
            use_field_extraction: If True, upload to /document/fieldextract to process text tags
            
        Returns:
            Document ID if successful, None otherwise
        """
        if not self.access_token:
            if not self.authenticate():
                return None
        try:
            with open(docx_path, 'rb') as f:
                content = f.read()
            
            files = {
                'file': (f"{document_name}.docx", content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            
            # Use field extraction endpoint to process text tags automatically
            if use_field_extraction:
                # Upload to /document/fieldextract to parse text tags and create fields
                # parse_type: "tag" - parse only simple tags (our format)
                data = { 'parse_type': 'tag' }
                endpoint = f"{self.base_url}/document/fieldextract"
            else:
                # Standard upload without field extraction
                data = { 'name': document_name }
                endpoint = f"{self.base_url}/document"
            
            headers = { "Authorization": f"Bearer {self.access_token}" }
            resp = requests.post(endpoint, files=files, data=data, headers=headers, timeout=180)
            resp.raise_for_status()
            doc_id = resp.json().get('id')
            
            return doc_id
        except requests.exceptions.RequestException as e:
            error_msg = str(e)
            if hasattr(e, 'response') and e.response is not None:
                try:
                    error_detail = e.response.json()
                    error_msg = json.dumps(error_detail, indent=2)
                except:
                    error_msg = e.response.text[:500]
            st.error(f"Failed to upload DOCX: {error_msg}")
            return None
        except Exception as e:
            st.error(f"Failed to upload DOCX: {str(e)}")
            return None

    def _find_signature_placeholders(self, docx_path: str) -> list:
        """
        Find SignNow Text Tags {{s_Signature_Signer1}} and {{s_Signature_Signer2}} in the merged DOCX file.
        Also supports legacy CLIENT_SIGNATURE and SALESMAN_SIGNATURE placeholders.
        
        Args:
            docx_path: Path to the merged DOCX file
            
        Returns:
            list: List of placeholder info dicts with 'type', 'paragraph_index', text, and position info
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            
            doc = Document(docx_path)
            placeholders = []
            
            # More accurate estimation: ~55 lines/paragraphs per page
            # Standard page is 792 points tall, average line is ~14pt
            lines_per_page = 55
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                # Check for Signer 1 text tag or legacy CLIENT_SIGNATURE
                if '{{s_Signature_Signer1}}' in text or 'CLIENT_SIGNATURE' in text.upper():
                    # Calculate estimated page number based on line position
                    estimated_page = para_idx // lines_per_page
                    
                    # Calculate Y position on the page
                    # Y coordinates: 0 is top, larger numbers go down
                    # Signature blocks are typically near the top (y=50-200)
                    line_on_page = para_idx % lines_per_page
                    # Each line is roughly 14 points
                    # Signatures are typically early on the page (lines 3-10)
                    y_position_on_page = 50 + (line_on_page * 14)  # Start at y=50, add 14pt per line
                    
                    para_format = paragraph.paragraph_format
                    space_before = para_format.space_before.pt if para_format.space_before else 0
                    space_after = para_format.space_after.pt if para_format.space_after else 0
                    
                    placeholder_type = '{{s_Signature_Signer1}}' if '{{s_Signature_Signer1}}' in text else 'CLIENT_SIGNATURE'
                    placeholders.append({
                        'type': placeholder_type,
                        'paragraph_index': para_idx,
                        'text': text,
                        'estimated_page': estimated_page,
                        'estimated_y': y_position_on_page,
                        'paragraph_format': {
                            'space_before': space_before,
                            'space_after': space_after,
                            'alignment': str(para_format.alignment) if para_format.alignment else None
                        }
                    })
                
                # Check for Signer 2 text tag or legacy SALESMAN_SIGNATURE
                elif '{{s_Signature_Signer2}}' in text or 'SALESMAN_SIGNATURE' in text.upper():
                    estimated_page = para_idx // lines_per_page
                    line_on_page = para_idx % lines_per_page
                    y_position_on_page = 50 + (line_on_page * 14)  # Start at y=50, add 14pt per line
                    
                    para_format = paragraph.paragraph_format
                    space_before = para_format.space_before.pt if para_format.space_before else 0
                    space_after = para_format.space_after.pt if para_format.space_after else 0
                    
                    placeholder_type = '{{s_Signature_Signer2}}' if '{{s_Signature_Signer2}}' in text else 'SALESMAN_SIGNATURE'
                    placeholders.append({
                        'type': placeholder_type,
                        'paragraph_index': para_idx,
                        'text': text,
                        'estimated_page': estimated_page,
                        'estimated_y': y_position_on_page,
                        'paragraph_format': {
                            'space_before': space_before,
                            'space_after': space_after,
                            'alignment': str(para_format.alignment) if para_format.alignment else None
                        }
                    })
            
            # Also search in tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if '{{s_Signature_Signer1}}' in cell_text or 'CLIENT_SIGNATURE' in cell_text.upper():
                            placeholder_type = '{{s_Signature_Signer1}}' if '{{s_Signature_Signer1}}' in cell_text else 'CLIENT_SIGNATURE'
                            placeholders.append({
                                'type': placeholder_type,
                                'table_index': table_idx,
                                'row_index': row_idx,
                                'col_index': col_idx,
                                'text': cell_text,
                                'estimated_page': 0,  # Tables are harder to estimate
                                'estimated_y': 0
                            })
                        elif '{{s_Signature_Signer2}}' in cell_text or 'SALESMAN_SIGNATURE' in cell_text.upper():
                            placeholder_type = '{{s_Signature_Signer2}}' if '{{s_Signature_Signer2}}' in cell_text else 'SALESMAN_SIGNATURE'
                            placeholders.append({
                                'type': placeholder_type,
                                'table_index': table_idx,
                                'row_index': row_idx,
                                'col_index': col_idx,
                                'text': cell_text,
                                'estimated_page': 0,
                                'estimated_y': 0
                            })
            
            return placeholders
            
        except Exception as e:
            st.warning(f"⚠️ Could not find signature placeholders: {str(e)}")
            import traceback
            st.code(traceback.format_exc())
            return []
    
    def create_and_send_merged_pair_docx(self, pair_type: str, contract_docx_path: str, terms_docx_path: str,
                                          document_name: str, email: str, salesman_email: str = None, 
                                          confirmed_by_name: str = None) -> tuple[bool, str]:
        """
        Merge two processed DOCX files into one DOCX, upload, and send for two-party signing.
        
        Args:
            pair_type: Type of document pair ('development_pair' or 'production_pair')
            contract_docx_path: Path to the contract DOCX file
            terms_docx_path: Path to the terms DOCX file
            document_name: Name for the merged document
            email: Email address of the prospect (signs first)
            salesman_email: Email address of the salesman (signs second). If None, uses account email
            confirmed_by_name: Name to populate in "Confirmed By" text field. If None, uses account username
            
        Returns:
            tuple[bool, str]: (success, message)
        """
        try:
            merged_docx = self._merge_docx(contract_docx_path, terms_docx_path)
            if not merged_docx or not os.path.exists(merged_docx):
                return False, "Failed to merge DOCX documents"
            
            # Upload using field extraction endpoint to automatically process text tags
            # This uploads to /document/fieldextract which processes text tags during upload
            doc_id = self.upload_docx(document_name, merged_docx, use_field_extraction=True)
            if not doc_id:
                return False, "Failed to upload merged DOCX with field extraction"
            
            # Wait a moment for SignNow to finish processing text tags (if needed)
            import time
            time.sleep(2)
            
            # Send for two-party signing (will verify fields exist from text tags)
            if self.send_document_for_two_party_signing(doc_id, email, salesman_email, confirmed_by_name):
                salesman_display = salesman_email or self.user_email or "salesman"
                return True, f"Merged document sent successfully. Prospect: {email}, Salesman: {salesman_display}. Document ID: {doc_id}"
            return False, "Merged DOCX uploaded but failed to send for signing"
        except Exception as e:
            return False, f"Error creating/sending merged DOCX: {str(e)}"
    
    def _docx_to_html(self, doc):
        """
        Convert docx Document to HTML with proper formatting preservation
        
        Args:
            doc: python-docx Document object
            
        Returns:
            str: HTML content
        """
        html_parts = []
        html_parts.append('<html><head><meta charset="UTF-8"></head><body>')
        
        # Check for images/logos first
        if hasattr(doc, 'inline_shapes') and doc.inline_shapes:
            html_parts.append('<div class="image-placeholder" style="color: blue; font-weight: bold; text-align: center; font-size: 16px;">IMAGE/LOGO NOT DISPLAYED IN PREVIEW</div>')
        
        for paragraph in doc.paragraphs:
            # Check if this paragraph contains an image (screenshot placeholder)
            has_image = False
            if paragraph.runs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        has_image = True
                        break
            
            # If paragraph has an image, show red placeholder
            if has_image:
                html_parts.append('<div class="screenshot-placeholder" style="color: red; font-weight: bold; text-align: center; font-size: 16px;">IMAGE ATTACHED GOES HERE</div>')
                continue
            
            if paragraph.text.strip():
                text = paragraph.text.strip()
                
                # Check for actual bold formatting in the document
                is_bold = paragraph.runs and any(run.bold for run in paragraph.runs)
                
                # Determine formatting based on content and actual formatting
                if any(word in text.upper() for word in ['DEVELOPMENT CONTRACT', 'TERMS AND CONDITIONS', 'PRODUCTION CONTRACT']):
                    html_parts.append(f'<h1>{text}</h1>')
                elif any(word in text.upper() for word in ['CONTRACT', 'AGREEMENT', 'TERMS', 'CONDITIONS']):
                    if len(text) < 100:  # Likely a heading
                        html_parts.append(f'<h2>{text}</h2>')
                    elif is_bold:
                        html_parts.append(f'<p class="bold">{text}</p>')
                    else:
                        html_parts.append(f'<p>{text}</p>')
                elif text.startswith('The following agreement'):
                    html_parts.append(f'<p class="bold">{text}</p>')
                elif 'total contract amount' in text.lower():
                    # Highlight contract amount
                    text_with_highlight = text.replace('$', '<span class="contract-amount">$').replace(' and is due', '</span> and is due')
                    html_parts.append(f'<p class="bold">{text_with_highlight}</p>')
                elif any(word in text.upper() for word in ['SIGNATURE', 'SIGNED', 'DATE', 'VITALINA', 'TEG INTL', 'ON BEHALF']):
                    html_parts.append(f'<p class="bold">{text}</p>')
                elif is_bold:
                    html_parts.append(f'<p class="bold">{text}</p>')
                else:
                    html_parts.append(f'<p>{text}</p>')
        
        # Add table placeholders
        if doc.tables:
            html_parts.append('<div class="table-placeholder">TABLE NOT DISPLAYED IN PREVIEW</div>')
        
        html_parts.append('</body></html>')
        return ''.join(html_parts)
    
    
    def send_document_for_signing(self, document_id: str, email: str, 
                                first_name: str, last_name: str) -> bool:
        """
        Send document for signing to the specified email
        
        Args:
            document_id: ID of the document to send
            email: Email address to send to
            first_name: First name of the signer
            last_name: Last name of the signer
            
        Returns:
            bool: True if successful, False otherwise
        """
        if not self.access_token:
            if not self.authenticate():
                return False
        
        try:
            send_url = f"{self.base_url}/document/{document_id}/invite"
            
            headers = {
                "Authorization": f"Bearer {self.access_token}",
                "Content-Type": "application/json"
            }
            
            data = {
                "to": email,
                "from": self.user_email
            }
            
            response = requests.post(send_url, json=data, headers=headers)
            response.raise_for_status()
            
            return True
            
        except requests.exceptions.RequestException as e:
            st.error(f"Failed to send document: {str(e)}")
            return False
        except Exception as e:
            st.error(f"Unexpected error sending document: {str(e)}")
            return False
    
    def _add_signature_fields_to_document(self, document_id: str, placeholders: list = None) -> bool:
        """
        Add signature fields to the document for two-party signing
        This must be done before sending the invite for role-based routing to work
        
        Args:
            document_id: ID of the document
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Get document info to determine page count and layout
            doc_url = f"{self.base_url}/document/{document_id}"
            headers = {
                "Authorization": f"Bearer {self.access_token}",
                "Content-Type": "application/json"
            }
            
            st.write("🔍 **DEBUG: Fetching document info...**")
            doc_response = requests.get(doc_url, headers=headers)
            doc_response.raise_for_status()
            doc_data = doc_response.json()
            st.write(f"✅ Document fetched. Document ID: {document_id}")
            st.json({"document_keys": list(doc_data.keys())})
            
            # Get page count
            pages = doc_data.get('pages', [])
            last_page_number = len(pages) - 1 if pages else 0
            st.write(f"📄 Document has {len(pages)} page(s).")
            
            # Get existing fields if any
            existing_fields = doc_data.get('fields', [])
            st.write(f"📝 Existing fields in document: {len(existing_fields)}")
            
            # Find text tags {{s_Signature_Signer1}} and {{s_Signature_Signer2}} in document
            # These should have been converted to fields by SignNow, but we check texts as fallback
            texts = doc_data.get('texts', [])
            signer1_positions = []
            signer2_positions = []
            
            # Method 1: Try to find text tags in SignNow's texts field (for debugging)
            if texts:
                st.write(f"📋 SignNow texts field contains {len(texts)} text elements")
            
            for text_item in texts:
                if isinstance(text_item, dict):
                    # Try different possible text field structures
                    text_content = (text_item.get('data', {}).get('text', '') or 
                                   text_item.get('text', '') or
                                   str(text_item.get('value', '')) or
                                   str(text_item))
                    
                    # Also check all keys/values in the dict
                    all_text_in_item = ' '.join([str(v) for v in text_item.values() if isinstance(v, (str, int, float))])
                    
                    # Look for SignNow text tags: {{s_Signature_Signer1}} or {{s_Signature_Signer2}}
                    if '{{s_Signature_Signer1}}' in text_content or 's_Signature_Signer1' in all_text_in_item:
                        signer1_positions.append({
                            'x': text_item.get('x', 0),
                            'y': text_item.get('y', 0),
                            'page_number': text_item.get('page_number', 0),
                            'source': 'signnow_api',
                            'text_item': text_item
                        })
                    elif '{{s_Signature_Signer2}}' in text_content or 's_Signature_Signer2' in all_text_in_item:
                        signer2_positions.append({
                            'x': text_item.get('x', 0),
                            'y': text_item.get('y', 0),
                            'page_number': text_item.get('page_number', 0),
                            'source': 'signnow_api',
                            'text_item': text_item
                        })
            
            # Method 2: Use DOCX-based placeholder positions from merged document analysis
            if placeholders:
                st.write(f"📄 Using {len(placeholders)} placeholder(s) from DOCX analysis")
                
                # Use the placeholder positions from the merged DOCX file
                # These are based on actual paragraph positions in the merged document
                for placeholder in placeholders:
                    para_idx = placeholder.get('paragraph_index', 0)
                    placeholder_type = placeholder.get('type', '')
                    
                    # Use estimated page and Y from DOCX analysis
                    # Ensure page number is within valid range
                    estimated_page = min(placeholder.get('estimated_page', 0), last_page_number)
                    estimated_page = max(0, estimated_page)  # Ensure non-negative
                    estimated_y = placeholder.get('estimated_y', 700)
                    
                    # X position: Signer 1 on left, Signer 2 on right  
                    # Based on image, signatures are near top, Signer 1 left, Signer 2 right
                    if placeholder_type in ['CLIENT_SIGNATURE', '{{s_Signature_Signer1}}', 's_Signature_Signer1']:
                        x_pos = 100  # Left side
                    else:  # SALESMAN_SIGNATURE or {{s_Signature_Signer2}}
                        x_pos = 350  # Right side
                    
                    # Y position: use the estimated Y from DOCX analysis
                    # Based on image, signatures are near TOP of page (y ~50-200)
                    # The estimated_y is already calculated from line position
                    y_pos = int(round(estimated_y))
                    
                    # Ensure coordinates are valid integers within page bounds
                    x_pos = int(round(x_pos))
                    x_pos = max(0, min(x_pos, 612))  # Page width is 612 points
                    y_pos = max(0, min(y_pos, 792))  # Page height is 792 points
                    estimated_page = int(round(estimated_page))
                    estimated_page = max(0, min(estimated_page, last_page_number))
                    
                    pos_dict = {
                        'x': x_pos,
                        'y': y_pos,
                        'page_number': estimated_page,
                        'source': 'docx_analysis',
                        'paragraph_index': para_idx,
                        'placeholder_text': placeholder.get('text', '')[:50]
                    }
                    
                    if placeholder_type in ['CLIENT_SIGNATURE', '{{s_Signature_Signer1}}', 's_Signature_Signer1']:
                        signer1_positions.append(pos_dict)
                    elif placeholder_type in ['SALESMAN_SIGNATURE', '{{s_Signature_Signer2}}', 's_Signature_Signer2']:
                        signer2_positions.append(pos_dict)
                
                st.write(f"✅ Mapped DOCX text tags to SignNow coordinates")
            
            st.write(f"📍 Found {len(signer1_positions)} Signer 1 text tag(s) ({{{{s_Signature_Signer1}}}})")
            st.write(f"📍 Found {len(signer2_positions)} Signer 2 text tag(s) ({{{{s_Signature_Signer2}}}})")
            
            if not signer1_positions or not signer2_positions:
                st.error("❌ Could not find signature text tags. Cannot proceed without exact positions.")
                return False
            
            # Add signature fields to the document
            # Coordinates are approximate - adjust based on your document layout
            # Standard page is 612x792 points (8.5" x 11" at 72 DPI)
            # Correct endpoint is PUT /document/{document_id} (not /document/{document_id}/field)
            fields_url = f"{self.base_url}/document/{document_id}"
            
            # Signature fields: 2 for client (Signer 1) and 2 for salesman (Signer 2)
            # Use placeholder positions if found, otherwise use estimated positions
            new_fields = []
            
            # Use found placeholder positions to place signature fields
            # Place signatures at the placeholder locations (slightly offset above the placeholder text)
            signature_offset_y = -40  # Place signature field above the placeholder text
            
            if signer1_positions and signer2_positions:
                # We found text tags - use their positions (should already be processed by SignNow)
                # Sort positions by page number to get contract (first) and terms (second) signatures
                signer1_sorted = sorted(signer1_positions, key=lambda p: p['page_number'])
                signer2_sorted = sorted(signer2_positions, key=lambda p: p['page_number'])
                
                # Create 2 signatures for each signer (assuming first 2 are for contract/terms)
                # Note: This should only run if SignNow didn't process text tags automatically
                for i, pos in enumerate(signer1_sorted[:2]):  # First 2 Signer 1 text tag positions
                    x_pos = int(pos.get('x', 100))
                    y_pos = int(pos.get('y', 150) + signature_offset_y)  # Above placeholder
                    page_num = int(pos.get('page_number', 0))
                    
                    # Ensure coordinates are within valid ranges
                    x_pos = max(0, min(x_pos, 612))  # Page width is 612 points
                    y_pos = max(0, min(y_pos, 792))  # Page height is 792 points
                    page_num = max(0, min(page_num, last_page_number))
                    
                    new_fields.append({
                        "x": x_pos,
                        "y": y_pos,
                        "width": 200,
                        "height": 25,  # Half of original height
                        "page_number": page_num,
                        "role": "Signer 1",
                        "required": True,
                        "type": "signature"
                    })
                
                for i, pos in enumerate(signer2_sorted[:2]):  # First 2 Signer 2 text tag positions
                    x_pos = int(pos.get('x', 350))
                    y_pos = int(pos.get('y', 150) + signature_offset_y)  # Above placeholder
                    page_num = int(pos.get('page_number', 0))
                    
                    # Ensure coordinates are within valid ranges
                    x_pos = max(0, min(x_pos, 612))  # Page width is 612 points
                    y_pos = max(0, min(y_pos, 792))  # Page height is 792 points
                    page_num = max(0, min(page_num, last_page_number))
                    
                    new_fields.append({
                        "x": x_pos,
                        "y": y_pos,
                        "width": 200,
                        "height": 25,  # Half of original height
                        "page_number": page_num,
                        "role": "Signer 2",
                        "required": True,
                        "type": "signature"
                    })
            else:
                # Fallback: Use estimated positions if placeholders not found
                st.warning("⚠️ Using estimated positions for signature fields")
                contract_signature_page = max(0, last_page_number - 2) if last_page_number >= 2 else 0
                terms_signature_page = last_page_number
                
                new_fields = [
                    # Signer 1 - Contract signature
                    {
                        "x": 100,
                        "y": 150,
                        "width": 200,
                        "height": 25,
                        "page_number": contract_signature_page,
                        "role": "Signer 1",
                        "required": True,
                        "type": "signature"
                    },
                    # Signer 1 - Terms signature
                    {
                        "x": 100,
                        "y": 150,
                        "width": 200,
                        "height": 25,
                        "page_number": terms_signature_page,
                        "role": "Signer 1",
                        "required": True,
                        "type": "signature"
                    },
                    # Signer 2 - Contract signature
                    {
                        "x": 350,
                        "y": 150,
                        "width": 200,
                        "height": 25,
                        "page_number": contract_signature_page,
                        "role": "Signer 2",
                        "required": True,
                        "type": "signature"
                    },
                    # Signer 2 - Terms signature
                    {
                        "x": 350,
                        "y": 150,
                        "width": 200,
                        "height": 25,
                        "page_number": terms_signature_page,
                        "role": "Signer 2",
                        "required": True,
                        "type": "signature"
                    }
                ]
            
            # Combine existing fields with new fields
            all_fields = existing_fields + new_fields
            st.write(f"📝 Total fields to add: {len(new_fields)}. Total fields after merge: {len(all_fields)}")
            
            # Show where signatures will be placed
            if new_fields:
                st.write("📍 **Signature field positions:**")
                for i, field in enumerate(new_fields, 1):
                    st.write(f"  Field {i}: Page {field['page_number']}, Role: {field['role']}, Position: ({field['x']}, {field['y']})")
            
            # Update document with fields
            # SignNow API expects the fields array in the update payload
            update_data = {
                "fields": all_fields
            }
            
            st.write("🔍 **DEBUG: Request details:**")
            st.write(f"URL: PUT {fields_url}")
            st.json({"update_data_structure": {
                "fields_count": len(update_data["fields"]),
                "first_field_sample": update_data["fields"][0] if update_data["fields"] else None
            }})
            
            field_response = requests.put(fields_url, json=update_data, headers=headers)
            
            st.write(f"📡 **DEBUG: Response status code: {field_response.status_code}**")
            
            # Better error handling to see what SignNow expects
            if field_response.status_code != 200:
                st.error(f"❌ Failed to add signature fields. Status: {field_response.status_code}")
                try:
                    error_detail = field_response.json()
                    st.error("**Error response:**")
                    st.json(error_detail)
                except:
                    st.error(f"**Error response text:** {field_response.text}")
                return False
            
            # Success
            try:
                response_data = field_response.json()
                st.success("✅ Signature fields added successfully!")
                st.json({"response_keys": list(response_data.keys()) if isinstance(response_data, dict) else "Non-dict response"})
            except:
                st.success("✅ Signature fields added successfully!")
            
            return True
            
        except requests.exceptions.RequestException as e:
            st.error(f"❌ Request exception: {str(e)}")
            # Show response if available
            if hasattr(e, 'response') and e.response is not None:
                st.error(f"**Response status:** {e.response.status_code}")
                try:
                    error_detail = e.response.json()
                    st.error("**Error details:**")
                    st.json(error_detail)
                except:
                    st.error(f"**Error response text:** {e.response.text}")
            return False
        except Exception as e:
            st.error(f"❌ Unexpected error: {str(e)}")
            import traceback
            st.code(traceback.format_exc())
            return False
    
    def send_document_for_two_party_signing(self, document_id: str, prospect_email: str,
                                           salesman_email: str = None, confirmed_by_name: str = None,
                                           placeholders: list = None) -> bool:
        """
        Send document for two-party signing with routing order:
        Step 1: Prospect signs first
        Step 2: Salesman signs second (confirmation)
        
        Args:
            document_id: ID of the document to send
            prospect_email: Email address of the prospect (signs first)
            salesman_email: Email address of the salesman (signs second). If None, uses self.user_email
            confirmed_by_name: Name to populate in "Confirmed By" text field. If None, uses account username
            
        Returns:
            bool: True if successful, False otherwise
        """
        if not self.access_token:
            if not self.authenticate():
                return False
        
        # Use salesman email from account if not provided
        if not salesman_email:
            salesman_email = self.user_email
        
        # Get salesman name from username if not provided
        if not confirmed_by_name and self.username:
            confirmed_by_name = self.username.split('@')[0]  # Use part before @ as name
        
        try:
            send_url = f"{self.base_url}/document/{document_id}/invite"
            
            headers = {
                "Authorization": f"Bearer {self.access_token}",
                "Content-Type": "application/json"
            }
            
            # Fetch document to check if fields were created from Text Tags
            doc_url = f"{self.base_url}/document/{document_id}"
            doc_response = requests.get(doc_url, headers=headers)
            doc_response.raise_for_status()
            doc_data = doc_response.json()
            
            # Check if document already has fields from Text Tags
            # SignNow Text Tags {{s_Signature_Signer1}} and {{s_Signature_Signer2}} should create fields automatically
            existing_fields = doc_data.get('fields', [])
            
            # Filter signature fields by role
            # Note: SignNow may assign roles based on text tag "o" parameter
            signer1_fields = [f for f in existing_fields 
                            if (f.get('role') == 'Signer 1' or f.get('role') == 'Signer1' or 
                                'Signer1' in str(f.get('name', '')) or 'Signer 1' in str(f.get('name', '')))
                            and f.get('type') == 'signature']
            
            signer2_fields = [f for f in existing_fields 
                            if (f.get('role') == 'Signer 2' or f.get('role') == 'Signer2' or 
                                'Signer2' in str(f.get('name', '')) or 'Signer 2' in str(f.get('name', '')))
                            and f.get('type') == 'signature']
            
            # Check if document already has fields from Text Tags
            # We need 2 fields for Signer 1 and 2 fields for Signer 2 (contract + terms)
            if len(signer1_fields) < 2 or len(signer2_fields) < 2:
                # No fields found - SignNow may not have processed text tags yet
                if len(existing_fields) == 0:
                    import time
                    time.sleep(3)
                    
                    # Check again
                    doc_response = requests.get(doc_url, headers=headers)
                    doc_response.raise_for_status()
                    doc_data = doc_response.json()
                    existing_fields = doc_data.get('fields', [])
                    signer1_fields = [f for f in existing_fields 
                                    if (f.get('role') == 'Signer 1' or f.get('role') == 'Signer1')
                                    and f.get('type') == 'signature']
                    signer2_fields = [f for f in existing_fields 
                                    if (f.get('role') == 'Signer 2' or f.get('role') == 'Signer2')
                                    and f.get('type') == 'signature']
                
                if len(signer1_fields) < 2 or len(signer2_fields) < 2:
                    st.error(f"❌ Text Tags did not create expected fields. Found {len(signer1_fields)} Signer 1 and {len(signer2_fields)} Signer 2 fields (need 2 each).")
                    return False
            
            # SignNow multi-signer routing now that fields are added
            # Note: subject and message fields removed - require subscription upgrade
            data = {
                "to": [
                    {
                        "email": prospect_email,
                        "role_id": "",
                        "role": "Signer 1",
                        "order": 1
                    },
                    {
                        "email": salesman_email,
                        "role_id": "",
                        "role": "Signer 2",
                        "order": 2
                    }
                ],
                "from": self.user_email
            }
            
            response = requests.post(send_url, json=data, headers=headers)
            
            # Better error handling
            if response.status_code != 200:
                try:
                    error_detail = response.json()
                    error_msg = json.dumps(error_detail, indent=2)
                    st.error(f"SignNow API Error ({response.status_code}): {error_msg}")
                except:
                    st.error(f"SignNow API Error ({response.status_code}): {response.text}")
                response.raise_for_status()
            else:
                response.raise_for_status()
            
            # Note: "Confirmed By" field should be included in document template with text tag for "Signer 2"
            # No need to map it programmatically - SignNow will handle it if the text tag exists
            
            return True
            
        except requests.exceptions.RequestException as e:
            st.error(f"Failed to send document for two-party signing: {str(e)}")
            return False
        except Exception as e:
            st.error(f"Unexpected error sending document for two-party signing: {str(e)}")
            return False
    
    def _add_confirmed_by_field_mapping(self, document_id: str, confirmed_by_name: str) -> bool:
        """
        Add or update field mapping for "Confirmed By" text field assigned to Signer 2 (salesman)
        
        This method attempts to add a text field for "Confirmed By" to the document.
        Note: Adding fields requires coordinates. If the document template already has a text field
        with role "Signer 2" or text tag {{s_Text_ConfirmedBy_Signer2}}, SignNow will handle it automatically.
        
        Args:
            document_id: ID of the document
            confirmed_by_name: Name to populate in the "Confirmed By" field
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Get document structure to find last page and suitable location
            doc_url = f"{self.base_url}/document/{document_id}"
            
            headers = {
                "Authorization": f"Bearer {self.access_token}",
                "Content-Type": "application/json"
            }
            
            # Get document info to find page count and field locations
            doc_response = requests.get(doc_url, headers=headers)
            doc_response.raise_for_status()
            document_data = doc_response.json()
            
            # Try to find existing signature fields to place "Confirmed By" near them
            # Default to last page, near bottom (typical signature location)
            # Coordinates are in points (1 point = 1/72 inch)
            # A4 page is approximately 595 x 842 points
            
            # For now, we'll attempt to add the field at a reasonable location
            # This is a best-effort approach - ideally the template should have the field
            fields_url = f"{self.base_url}/document/{document_id}/field"
            
            # Estimate coordinates based on standard signature area (bottom of last page)
            # This is approximate and may need adjustment based on your document layout
            field_data = {
                "fields": [
                    {
                        "x": 100,  # Left margin
                        "y": 100,  # Near bottom (adjust based on your document)
                        "width": 200,
                        "height": 30,
                        "page_number": 0,  # Last page (0-indexed, adjust based on document)
                        "label": "Confirmed By",
                        "role": "Signer 2",
                        "required": False,
                        "type": "text",
                        "prefilled_text": confirmed_by_name
                    }
                ]
            }
            
            # Attempt to add the field
            field_response = requests.put(fields_url, json=field_data, headers=headers)
            
            if field_response.status_code in [200, 201]:
                return True
            else:
                # Field addition failed - likely because coordinates need adjustment or template should have it
                # This is not critical - the two-party signing will still work
                return False
            
        except requests.exceptions.RequestException as e:
            # Field mapping is optional - log but don't fail
            return False
        except Exception as e:
            # Field mapping is optional - log but don't fail
            return False
    
    def create_and_send_contract(self, template_type: str, client_name: str, email: str, 
                               contract_amount: str = None, contract_date: str = None) -> Tuple[bool, str]:
        """
        Complete workflow: create document from template and send for signing
        
        Args:
            template_type: Type of template to use (development_contract, development_terms, terms_conditions, production_contract)
            client_name: Client name
            email: Email address of the signer
            contract_amount: Contract amount (optional, required for contract types)
            contract_date: Contract date (optional, defaults to current date)
            
        Returns:
            Tuple[bool, str]: (success, message)
        """
        try:
            # Set default date if not provided
            if not contract_date:
                from datetime import datetime
                contract_date = datetime.now().strftime("%B %d, %Y")
            
            document_name = f"{template_type}_{client_name.replace(' ', '_')}_{contract_date.replace(' ', '_').replace(',', '')}"
            
            # Create document
            document_id = self.create_document_from_template(
                template_type, document_name, client_name, email, contract_amount, contract_date
            )
            
            if not document_id:
                return False, "Failed to create document from template"
            
            # Send for signing
            if self.send_document_for_signing(document_id, email, client_name, ""):
                return True, f"Contract sent successfully to {email}. Document ID: {document_id}"
            else:
                return False, "Document created but failed to send for signing"
                
        except Exception as e:
            return False, f"Error in contract workflow: {str(e)}"
    
    def create_and_send_document_pair(self, pair_type, client_name, email, 
                                    contract_amount=None, contract_date=None,
                                    deposit_amount=None, total_contract_amount=None,
                                    sewing_cost=None, pre_production_fee=None,
                                    tegmade_for=None):
        """
        Create and send a document pair for signing
        
        Args:
            pair_type: Type of document pair ('development_pair' or 'production_pair')
            client_name: Client name
            email: Email address of the signer
            contract_amount: Contract amount (for development contracts)
            contract_date: Contract date (optional, defaults to current date)
            deposit_amount: Deposit amount (for production contracts)
            total_contract_amount: Total contract amount (for production contracts)
            sewing_cost: Sewing cost (for production contracts)
            pre_production_fee: Pre-production fee (for production contracts)
            tegmade_for: Name to replace VITALINA GHINZELLI (optional)
            
        Returns:
            Tuple[bool, str]: (success, message)
        """
        try:
            # Set default date if not provided
            if not contract_date:
                from datetime import datetime
                contract_date = datetime.now().strftime("%B %d, %Y")
            
            # Determine which documents to process
            if pair_type == 'development_pair':
                # Process development contract and terms
                contract_template = 'development_contract'
                terms_template = 'development_terms'
            elif pair_type == 'production_pair':
                # Process production contract and terms
                contract_template = 'production_contract'
                terms_template = 'production_terms'
            else:
                return False, f"Unknown pair type: {pair_type}"
            
            # Process both documents
            from docx_template_processor import DocxTemplateProcessor
            processor = DocxTemplateProcessor()
            
            # Process contract document
            contract_path = processor.process_document(
                template_type=contract_template,
                client_name=client_name,
                email=email,
                contract_amount=contract_amount,
                contract_date=contract_date,
                deposit_amount=deposit_amount,
                total_contract_amount=total_contract_amount,
                sewing_cost=sewing_cost,
                pre_production_fee=pre_production_fee,
                tegmade_for=tegmade_for
            )
            
            # Process terms document
            terms_path = processor.process_document(
                template_type=terms_template,
                client_name=client_name,
                email=email,
                contract_amount=None,
                contract_date=contract_date,
                tegmade_for=tegmade_for
            )
            
            # Upload both documents to SignNow
            document_ids = []
            
            for doc_path, doc_type in [(contract_path, "Contract"), (terms_path, "Terms")]:
                with open(doc_path, 'rb') as docx_file:
                    docx_content = docx_file.read()
                
                actual_filename = os.path.basename(doc_path)
                files = {
                    'file': (actual_filename, docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                }
                
                data = {
                    'name': f'{pair_type}_{doc_type} - {client_name}'
                }
                
                headers = {
                    "Authorization": f"Bearer {self.access_token}"
                }
                
                # Create document in SignNow
                response = requests.post(
                    f"{self.base_url}/document",
                    files=files,
                    data=data,
                    headers=headers,
                    timeout=90
                )
                response.raise_for_status()
                
                create_response = response.json()
                document_id = create_response.get("id")
                document_ids.append(document_id)
                
                # Clean up processed document
                os.remove(doc_path)
            
            # Send both documents for signing
            for document_id in document_ids:
                send_url = f"{self.base_url}/document/{document_id}/invite"
                send_data = {
                    "to": email,
                    "from": self.user_email
                }
                
                send_response = requests.post(send_url, json=send_data, headers=headers)
                send_response.raise_for_status()
            
            return True, f"Document pair sent successfully to {email}. Document IDs: {', '.join(document_ids)}"
                
        except Exception as e:
            return False, f"Error creating and sending document pair: {str(e)}"


def load_signnow_credentials(account_name: str = None) -> Dict[str, str]:
    """
    Load SignNow credentials from Streamlit secrets
    
    Args:
        account_name: Name of the account to load ('heather', 'jennifer', 'anthony', or None for default)
    
    Returns:
        Dict containing SignNow credentials
    """
    try:
        # Determine which secrets section to use
        if account_name:
            # Use specific account section
            secrets_key = f'signnow - {account_name.lower()}'
        else:
            # Fallback to default signnow section if no account specified
            secrets_key = 'signnow'
        
        if secrets_key not in st.secrets:
            if account_name:
                st.error(f"SignNow configuration for '{account_name}' not found in secrets.toml")
            else:
                st.error("SignNow configuration not found in secrets.toml")
            return {}
        
        signnow_config = st.secrets[secrets_key]
        
        required_fields = ['client_id', 'client_secret', 'basic_auth_token', 'username', 'password']
        for field in required_fields:
            if field not in signnow_config:
                st.error(f"SignNow {field} not found in secrets.toml for account '{account_name}'")
                return {}
        
        return signnow_config
        
    except Exception as e:
        st.error(f"Error reading SignNow secrets: {str(e)}")
        return {}