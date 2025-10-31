#!/usr/bin/env python3
"""
Scan DOCX files to identify placeholders that need to be replaced
"""

import os
from docx import Document
import re
from collections import defaultdict

def scan_docx_for_placeholders(file_path):
    """Scan a DOCX file and identify potential placeholders"""
    doc = Document(file_path)
    
    # Common patterns that might be placeholders
    placeholder_patterns = {
        'dates': r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}\b',
        'amounts': r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?',
        'names': r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+',  # Proper names
    }
    
    all_texts = []
    unique_texts = set()
    
    # Scan paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            all_texts.append(('paragraph', text))
            unique_texts.add(text)
    
    # Scan tables
    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        all_texts.append((f'table_{table_idx}', text))
                        unique_texts.add(text)
    
    # Find potential placeholders
    potential_placeholders = {
        'dates': [],
        'amounts': [],
        'names': [],
        'uppercase_placeholders': [],  # Like "CLIENT NAME", "BRAND NAME"
        'vitalina_patterns': []  # VITALINA GHINZELLI or similar
    }
    
    for text in unique_texts:
        # Check for dates
        dates = re.findall(placeholder_patterns['dates'], text)
        if dates:
            potential_placeholders['dates'].extend(dates)
        
        # Check for amounts
        amounts = re.findall(placeholder_patterns['amounts'], text)
        if amounts:
            potential_placeholders['amounts'].extend(amounts)
        
        # Check for uppercase placeholder patterns (like "CLIENT NAME", "BRAND NAME")
        uppercase_pattern = r'\b[A-Z]{2,}(?:\s+[A-Z]{2,})+\b'
        uppercase_matches = re.findall(uppercase_pattern, text)
        for match in uppercase_matches:
            # Filter out common words
            if match not in ['TEG', 'TEG INT', 'DATE', 'TEG INT L', 'AND', 'FOR', 'THE', 'ALL', 'ARE', 'BASED', 'ON', 'GIVEN', 'SPECIFICATION', 'SHEETS']:
                potential_placeholders['uppercase_placeholders'].append(match)
        
        # Check for standalone "NAME" placeholder (not part of "CLIENT NAME" or "BRAND NAME")
        if re.search(r'\bNAME\b', text) and 'CLIENT NAME' not in text and 'BRAND NAME' not in text:
            # Check if it's a standalone placeholder pattern like "NAME and on behalf of" or similar
            if re.search(r'\bNAME\s+(and|&)', text, re.IGNORECASE) or re.search(r'NAME\s*:', text):
                potential_placeholders['uppercase_placeholders'].append('NAME')
        
        # Check for "DATE" or "Date:" placeholder patterns
        if re.search(r'\bDATE\b|Date:\s*(TBD|NAME|\b\w+\b)', text, re.IGNORECASE):
            potential_placeholders['uppercase_placeholders'].append('DATE')
        
        # Check for VITALINA or similar patterns
        if 'VITALINA' in text.upper() or 'GHINZELLI' in text.upper() or 'TEGMADE' in text.upper():
            # Extract the full name
            words = text.split()
            for i, word in enumerate(words):
                if 'VITALINA' in word.upper():
                    if i + 1 < len(words):
                        full_name = f"{word} {words[i+1]}"
                        potential_placeholders['vitalina_patterns'].append(full_name)
                    break
        
        # Check for names (excluding common words)
        words = text.split()
        for word in words:
            if re.match(placeholder_patterns['names'], word):
                # Skip common words that might match
                if word not in ['Development', 'Production', 'Terms', 'Conditions', 'Contract', 'Agreement']:
                    potential_placeholders['names'].append(word)
    
    # Get unique values
    unique_dates = sorted(set(potential_placeholders['dates']))
    unique_amounts = sorted(set(potential_placeholders['amounts']))
    unique_names = sorted(set(potential_placeholders['names']))
    unique_uppercase = sorted(set(potential_placeholders['uppercase_placeholders']))
    unique_vitalina = sorted(set(potential_placeholders['vitalina_patterns']))
    
    return {
        'file': os.path.basename(file_path),
        'total_paragraphs': len([t for t in all_texts if t[0] == 'paragraph']),
        'total_tables': len(doc.tables),
        'dates': unique_dates,
        'amounts': unique_amounts,
        'names': unique_names[:10],  # Limit to first 10 names
        'uppercase_placeholders': unique_uppercase,
        'vitalina_patterns': unique_vitalina,
        'sample_texts': list(unique_texts)[:30]  # First 30 unique texts
    }

def main():
    """Scan all DOCX files in inputs folder"""
    inputs_folder = 'inputs'
    
    docx_files = [
        'Development Contract.docx',
        'Development Terms and Conditions .docx',
        'Production Contract.docx',
        'Production Terms and Conditions.docx'
    ]
    
    print("=" * 80)
    print("Scanning DOCX files for placeholders")
    print("=" * 80)
    
    results = {}
    
    for docx_file in docx_files:
        file_path = os.path.join(inputs_folder, docx_file)
        
        if not os.path.exists(file_path):
            print(f"\n❌ File not found: {file_path}")
            continue
        
        print(f"\n📄 Scanning: {docx_file}")
        print("-" * 80)
        
        try:
            result = scan_docx_for_placeholders(file_path)
            results[docx_file] = result
            
            print(f"Total paragraphs: {result['total_paragraphs']}")
            print(f"Total tables: {result['total_tables']}")
            
            if result['dates']:
                print(f"\n📅 Dates found ({len(result['dates'])}):")
                for date in result['dates']:
                    print(f"   - {date}")
            
            if result['amounts']:
                print(f"\n💰 Amounts found ({len(result['amounts'])}):")
                for amount in result['amounts']:
                    print(f"   - {amount}")
            
            if result['names']:
                print(f"\n👤 Names found ({len(result['names'])}):")
                for name in result['names']:
                    print(f"   - {name}")
            
            if result['uppercase_placeholders']:
                print(f"\n🔤 UPPERCASE Placeholders found ({len(result['uppercase_placeholders'])}):")
                for placeholder in result['uppercase_placeholders']:
                    print(f"   - {placeholder}")
            
            if result['vitalina_patterns']:
                print(f"\n✍️ Signature name patterns found ({len(result['vitalina_patterns'])}):")
                for pattern in result['vitalina_patterns']:
                    print(f"   - {pattern}")
            
            print(f"\n📝 Sample texts (first 10):")
            for i, text in enumerate(result['sample_texts'][:10], 1):
                print(f"   {i}. {text[:100]}...")
                
        except Exception as e:
            print(f"❌ Error scanning {docx_file}: {str(e)}")
    
    # Generate suggested template mapping
    print("\n" + "=" * 80)
    print("SUGGESTED TEMPLATE MAPPING")
    print("=" * 80)
    
    template_key_map = {
        'Development Contract.docx': 'development_contract',
        'Development Terms and Conditions .docx': 'development_terms',
        'Production Contract.docx': 'production_contract',
        'Production Terms and Conditions.docx': 'production_terms'
    }
    
    for docx_file, template_key in template_key_map.items():
        if docx_file in results:
            result = results[docx_file]
            print(f"\n# {template_key}")
            print("'replacements': {")
            
            # Add dates
            if result['dates']:
                print(f"    '{result['dates'][0]}': 'CONTRACT_DATE',")
            
            # Add amounts
            if result['amounts']:
                if template_key == 'development_contract':
                    if result['amounts']:
                        print(f"    '{result['amounts'][0]}': 'CONTRACT_AMOUNT',")
                elif template_key == 'production_contract':
                    for i, amount in enumerate(result['amounts'][:4]):
                        if i == 0:
                            print(f"    '{amount}': 'TOTAL_CONTRACT_AMOUNT',")
                        elif i == 1:
                            print(f"    '{amount}': 'SEWING_COST',")
                        elif i == 2:
                            print(f"    '{amount}': 'PRE_PRODUCTION_FEE',")
                        elif i == 3:
                            print(f"    '{amount}': 'TOTAL_DUE_AT_SIGNING',")
            
            # Add uppercase placeholders
            if result['uppercase_placeholders']:
                for placeholder in result['uppercase_placeholders']:
                    if 'CLIENT' in placeholder and 'NAME' in placeholder:
                        print(f"    '{placeholder}': 'CLIENT_NAME',")
                    elif 'BRAND' in placeholder and 'NAME' in placeholder:
                        print(f"    '{placeholder}': 'TEGMADE_FOR',")
            
            # Add names (look for VITALINA GHINZELLI or similar)
            if result['vitalina_patterns']:
                for pattern in result['vitalina_patterns']:
                    print(f"    '{pattern}': 'TEGMADE_FOR',")
            
            print("},")
            
            # Check for multiple replacements (client name appears twice)
            client_name_found = False
            for text in result['sample_texts']:
                # Look for common client name patterns
                if any(word in text for word in result['names'][:3] if len(word) > 3):
                    print(f"# 'multiple_replacements': {{")
                    print(f"#     'CLIENT_NAME_PLACEHOLDER': ['CLIENT_NAME', 'TEGMADE_FOR']")
                    print(f"# }}")
                    break

if __name__ == "__main__":
    main()

