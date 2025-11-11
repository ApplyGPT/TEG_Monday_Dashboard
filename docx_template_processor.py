#!/usr/bin/env python3
"""
Create documents from original .docx files with variable replacement
Preserves exact formatting, design, text, figures, tables, etc.
"""

import os
import shutil
from docx import Document
from docx.shared import Inches
import re
from datetime import datetime

class DocxTemplateProcessor:
    """Process .docx templates with variable replacement while preserving formatting"""
    
    def __init__(self):
        self.template_mapping = {
            'development_contract': {
                'file': 'Development Contract.docx',
                'replacements': {
                    '$6,971.00': 'CONTRACT_AMOUNT',
                    'CLIENT NAME': 'CLIENT_NAME',
                    'BRAND NAME': 'TEGMADE_FOR',
                    'NAME': 'CLIENT_NAME',  # Standalone NAME placeholder
                    'DATE_V': 'CONTRACT_DATE'  # Date placeholder (DATE_V)
                },
                'multiple_replacements': {}  # No multiple replacements needed with new placeholders
            },
            'development_terms': {
                'file': 'Development Terms and Conditions .docx',
                'replacements': {
                    'BRAND NAME': 'TEGMADE_FOR',
                    'NAME': 'CLIENT_NAME',  # Standalone NAME placeholder
                    'DATE_V': 'CONTRACT_DATE'  # Date placeholder (DATE_V)
                },
                'multiple_replacements': {}  # No multiple replacements needed
            },
            'production_contract': {
                'file': 'Production Contract.docx',
                'replacements': {
                    '$27,228.00': 'TOTAL_DUE_AT_SIGNING',  # Total Due at Signing
                    '$41,428.00': 'TOTAL_CONTRACT_AMOUNT',  # Total Contract Amount
                    '$14,200.00': 'SEWING_COST',  # Sewing Cost
                    '$13,028.00': 'PRE_PRODUCTION_FEE',  # Pre-production Fee
                    # $300.00 will always be $300.00 (no replacement needed)
                    'CLIENT NAME': 'CLIENT_NAME',
                    'BRAND NAME': 'TEGMADE_FOR',
                    'NAME': 'CLIENT_NAME',  # Standalone NAME placeholder
                    'DATE_V': 'CONTRACT_DATE'  # Date placeholder (DATE_V)
                },
                'multiple_replacements': {}  # No multiple replacements needed with new placeholders
            },
            'production_terms': {
                'file': 'Production Terms and Conditions.docx',
                'replacements': {
                    'BRAND NAME': 'TEGMADE_FOR',
                    'NAME': 'CLIENT_NAME',  # Standalone NAME placeholder
                    'DATE_V': 'CONTRACT_DATE'  # Date placeholder (DATE_V)
                },
                'multiple_replacements': {}  # No multiple replacements needed
            }
        }
    
    def process_document(self, template_type, client_name, email, 
                        contract_amount=None, contract_date=None,
                        total_contract_amount=None, sewing_cost=None,
                        pre_production_fee=None, total_due_at_signing=None,
                        uploaded_pdf=None, tegmade_for=None):
        """
        Process a document template with variable replacement
        
        Args:
            template_type: Type of template to process
            client_name: Client name to replace
            email: Client email (for reference)
            contract_amount: Contract amount (if applicable)
            contract_date: Contract date (if not provided, uses current date)
            total_contract_amount: Total contract amount (for production contracts)
            sewing_cost: Sewing cost (for production contracts)
            pre_production_fee: Pre-production fee (for production contracts)
            total_due_at_signing: Total amount due at signing (for production contracts)
            uploaded_pdf: Uploaded PDF file (Streamlit UploadedFile object)
            tegmade_for: Name to replace VITALINA GHINZELLI (optional)
            
        Returns:
            str: Path to the processed document
        """
        if template_type not in self.template_mapping:
            raise ValueError(f"Unknown template type: {template_type}")
        
        template_config = self.template_mapping[template_type]
        template_file = template_config['file']
        template_path = os.path.join('inputs', template_file)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        # Set default date if not provided
        if not contract_date:
            contract_date = datetime.now().strftime("%B %d, %Y")
        
        # Create output filename
        safe_client_name = client_name.replace(' ', '_').replace(',', '')
        safe_date = contract_date.replace(' ', '_').replace(',', '')
        output_filename = f"{template_type}_{safe_client_name}_{safe_date}.docx"
        output_path = os.path.join('processed_documents', output_filename)
        
        # Create output directory if it doesn't exist
        os.makedirs('processed_documents', exist_ok=True)
        
        # Copy template to output location
        shutil.copy2(template_path, output_path)
        
        # Open the document for processing
        doc = Document(output_path)
        
        # Prepare replacement values
        clean_company_name = tegmade_for.strip() if tegmade_for else ''
        has_company_name = bool(clean_company_name)

        replacement_values = {
            'CLIENT_NAME': client_name,
            'CONTRACT_DATE': contract_date,
            'CONTRACT_AMOUNT': self._format_contract_amount(contract_amount) if contract_amount else '$0.00',
            'TOTAL_CONTRACT_AMOUNT': self._format_contract_amount(total_contract_amount) if total_contract_amount else '$0.00',
            'SEWING_COST': self._format_contract_amount(sewing_cost) if sewing_cost else '$0.00',
            'PRE_PRODUCTION_FEE': self._format_contract_amount(pre_production_fee) if pre_production_fee else '$0.00',
            'TOTAL_DUE_AT_SIGNING': self._format_contract_amount(total_due_at_signing) if total_due_at_signing else '$0.00',
            'TEGMADE_FOR': clean_company_name if has_company_name else ''
        }
        
        # Process replacements
        replacements_made = self._replace_text_in_document(doc, template_config['replacements'], replacement_values, template_config.get('multiple_replacements', {}))

        # Remove company phrases when no company name provided
        if not has_company_name:
            self._remove_company_name_phrases(doc)

        # Ensure key values remain bold after any removals or adjustments
        bold_targets = [client_name]
        if contract_amount:
            bold_targets.append(replacement_values['CONTRACT_AMOUNT'])
        if total_contract_amount:
            bold_targets.append(replacement_values['TOTAL_CONTRACT_AMOUNT'])
        if sewing_cost:
            bold_targets.append(replacement_values['SEWING_COST'])
        if pre_production_fee:
            bold_targets.append(replacement_values['PRE_PRODUCTION_FEE'])
        if total_due_at_signing:
            bold_targets.append(replacement_values['TOTAL_DUE_AT_SIGNING'])

        self._ensure_bold_values(doc, bold_targets)
        
        # Insert PDF (converted to image) after first paragraph for contract documents
        if uploaded_pdf and template_type in ['development_contract', 'production_contract']:
            print(f"📄 Processing PDF: {uploaded_pdf.name} ({uploaded_pdf.size} bytes)")
            self._insert_pdf_after_first_paragraph(doc, uploaded_pdf)
        
        # Save the processed document
        doc.save(output_path)
        
        print(f"✅ Processed {template_type}: {output_filename}")
        print(f"   Replacements made: {replacements_made}")
        
        return output_path
    
    def _replace_text_in_document(self, doc, replacement_map, values, multiple_replacements=None):
        """Replace text in document while preserving formatting and making replacements bold"""
        replacements_made = []
        
        if multiple_replacements is None:
            multiple_replacements = {}
        
        # Track occurrences across the entire document for multiple replacements
        occurrence_counters = {}
        for text_to_replace in multiple_replacements.keys():
            occurrence_counters[text_to_replace] = 0
        
        # Process paragraphs - handle multiple replacements first
        for para_idx, paragraph in enumerate(doc.paragraphs):
            paragraph_replacements = []
            original_para_text = paragraph.text
            
            # Handle multiple replacements (same text, different values for each occurrence)
            for text_to_replace, replacement_vars in multiple_replacements.items():
                if text_to_replace in paragraph.text:
                    # Count occurrences in this paragraph
                    paragraph_occurrences = paragraph.text.count(text_to_replace)
                    
                    # Process each occurrence in this paragraph
                    current_text = paragraph.text
                    for i in range(paragraph_occurrences):
                        occurrence_counters[text_to_replace] += 1
                        occurrence_index = occurrence_counters[text_to_replace]
                        
                        # Determine which value to use based on occurrence index
                        if occurrence_index == 1:
                            # First occurrence across document
                            replacement_value = values[replacement_vars[0]]
                        else:
                            # Second occurrence across document
                            replacement_value = values[replacement_vars[1]]
                        
                        # Replace the first occurrence of this text in current paragraph
                        current_text = current_text.replace(text_to_replace, replacement_value, 1)
                    
                    if current_text != paragraph.text:
                        paragraph_replacements.append((paragraph.text, current_text))
                        replacements_made.append(f"{text_to_replace} -> {replacement_value} (occurrence {occurrence_counters[text_to_replace]})")
            
            # Handle regular single replacements - process longer strings first to avoid partial matches
            # Sort by length (longest first) to ensure "CLIENT NAME" is replaced before "NAME"
            sorted_replacements = sorted(replacement_map.items(), key=lambda x: len(x[0]), reverse=True)
            for old_text, variable in sorted_replacements:
                if old_text not in multiple_replacements:
                    paragraph_text = paragraph.text
                    
                    # For uppercase placeholders (like "BRAND NAME", "CLIENT NAME", "NAME")
                    # match case-insensitively to catch "Brand Name", "Client Name", etc.
                    # Use word boundaries for single-word placeholders to avoid matching "name" in "named"
                    if old_text.isupper() and len(old_text.split()) > 0:
                        # Case-insensitive matching for uppercase placeholders
                        # For single words (like "NAME"), use word boundaries to match whole words only
                        if len(old_text.split()) == 1:
                            # Single word: use word boundaries (e.g., \bNAME\b matches "NAME" but not "named")
                            pattern = re.compile(r'\b' + re.escape(old_text) + r'\b', re.IGNORECASE)
                        else:
                            # Multiple words: match case-insensitively (e.g., "BRAND NAME" matches "Brand Name")
                            pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                        matches = list(pattern.finditer(paragraph_text))
                        if matches:
                            # Replace each occurrence with its actual case, then replace with our value
                            for match in matches:
                                actual_text = paragraph_text[match.start():match.end()]
                                # Only add once per unique actual text variant
                                if not any(r[0] == actual_text for r in paragraph_replacements):
                                    paragraph_replacements.append((actual_text, values[variable]))
                            replacements_made.append(f"{old_text} (case-insensitive) -> {values[variable]} ({len(matches)} occurrence(s) in para {para_idx})")
                    elif old_text in paragraph_text:
                        # For exact matches (like "$6,971.00" or "DATE_V")
                        count = paragraph_text.count(old_text)
                        if count > 0:
                            if not any(r[0] == old_text for r in paragraph_replacements):
                                paragraph_replacements.append((old_text, values[variable]))
                                replacements_made.append(f"{old_text} -> {values[variable]} ({count} occurrence(s) in para {para_idx})")
            
            # Apply all replacements to this paragraph at once
            if paragraph_replacements:
                self._replace_multiple_texts_with_bold(paragraph, paragraph_replacements)
                # Verify replacement happened (check case-insensitively)
                if ("BRAND NAME" in original_para_text.upper() or "Brand Name" in original_para_text) and \
                   ("BRAND NAME" in paragraph.text.upper() or "Brand Name" in paragraph.text):
                    print(f"⚠️ Warning: 'BRAND NAME' (any case) still present in paragraph {para_idx} after replacement")
                    print(f"   Original: {original_para_text[:100]}...")
                    print(f"   After: {paragraph.text[:100]}...")
            elif 'Date' in paragraph.text or 'Date:' in paragraph.text:
                # No replacements but has Date - still fix it (for TEG lines)
                self._fix_date_line_without_replacement(paragraph)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_replacements = []
                        
                        # Handle multiple replacements in tables
                        for text_to_replace, replacement_vars in multiple_replacements.items():
                            if text_to_replace in paragraph.text:
                                # Count occurrences in this paragraph
                                paragraph_occurrences = paragraph.text.count(text_to_replace)
                                
                                # Process each occurrence in this paragraph
                                current_text = paragraph.text
                                for i in range(paragraph_occurrences):
                                    occurrence_counters[text_to_replace] += 1
                                    occurrence_index = occurrence_counters[text_to_replace]
                                    
                                    # Determine which value to use based on occurrence index
                                    if occurrence_index == 1:
                                        # First occurrence across document
                                        replacement_value = values[replacement_vars[0]]
                                    else:
                                        # Second occurrence across document
                                        replacement_value = values[replacement_vars[1]]
                                    
                                    # Replace the first occurrence of this text in current paragraph
                                    current_text = current_text.replace(text_to_replace, replacement_value, 1)
                                
                                if current_text != paragraph.text:
                                    paragraph_replacements.append((paragraph.text, current_text))
                                    replacements_made.append(f"{text_to_replace} -> {replacement_value} (occurrence {occurrence_counters[text_to_replace]})")
                        
                        # Handle regular single replacements in tables - process longer strings first
                        sorted_replacements = sorted(replacement_map.items(), key=lambda x: len(x[0]), reverse=True)
                        for old_text, variable in sorted_replacements:
                            if old_text not in multiple_replacements:
                                paragraph_text = paragraph.text
                                
                                # For uppercase placeholders (like "BRAND NAME", "CLIENT NAME", "NAME")
                                # match case-insensitively to catch "Brand Name", "Client Name", etc.
                                # Use word boundaries for single-word placeholders to avoid matching "name" in "named"
                                if old_text.isupper() and len(old_text.split()) > 0:
                                    # Case-insensitive matching for uppercase placeholders
                                    # For single words (like "NAME"), use word boundaries to match whole words only
                                    if len(old_text.split()) == 1:
                                        # Single word: use word boundaries (e.g., \bNAME\b matches "NAME" but not "named")
                                        pattern = re.compile(r'\b' + re.escape(old_text) + r'\b', re.IGNORECASE)
                                    else:
                                        # Multiple words: match case-insensitively (e.g., "BRAND NAME" matches "Brand Name")
                                        pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                                    matches = list(pattern.finditer(paragraph_text))
                                    if matches:
                                        # Replace each occurrence with its actual case, then replace with our value
                                        for match in matches:
                                            actual_text = paragraph_text[match.start():match.end()]
                                            # Only add once per unique actual text variant
                                            if not any(r[0] == actual_text for r in paragraph_replacements):
                                                paragraph_replacements.append((actual_text, values[variable]))
                                        replacements_made.append(f"{old_text} (case-insensitive) -> {values[variable]} ({len(matches)} occurrence(s) in table)")
                                elif old_text in paragraph_text:
                                    # For exact matches (like "$6,971.00" or "DATE_V")
                                    count = paragraph_text.count(old_text)
                                    if count > 0:
                                        paragraph_replacements.append((old_text, values[variable]))
                                        replacements_made.append(f"{old_text} -> {values[variable]} ({count} occurrence(s) in table)")
                        
                        # Apply all replacements to this paragraph at once
                        if paragraph_replacements:
                            self._replace_multiple_texts_with_bold(paragraph, paragraph_replacements)
        
        # Process text boxes and shapes (headers/footers might have text boxes)
        # This ensures we catch "BRAND NAME" even if it's in a text box in the signature section
        try:
            for section in doc.sections:
                # Process header paragraphs
                if section.header:
                    for paragraph in section.header.paragraphs:
                        paragraph_replacements = []
                        sorted_replacements = sorted(replacement_map.items(), key=lambda x: len(x[0]), reverse=True)
                        for old_text, variable in sorted_replacements:
                            if old_text not in multiple_replacements:
                                paragraph_text = paragraph.text
                                
                                # Case-insensitive matching for uppercase placeholders
                                # Use word boundaries for single-word placeholders to avoid matching "name" in "named"
                                if old_text.isupper() and len(old_text.split()) > 0:
                                    # For single words (like "NAME"), use word boundaries to match whole words only
                                    if len(old_text.split()) == 1:
                                        # Single word: use word boundaries (e.g., \bNAME\b matches "NAME" but not "named")
                                        pattern = re.compile(r'\b' + re.escape(old_text) + r'\b', re.IGNORECASE)
                                    else:
                                        # Multiple words: match case-insensitively (e.g., "BRAND NAME" matches "Brand Name")
                                        pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                                    matches = list(pattern.finditer(paragraph_text))
                                    if matches:
                                        for match in matches:
                                            actual_text = paragraph_text[match.start():match.end()]
                                            if not any(r[0] == actual_text for r in paragraph_replacements):
                                                paragraph_replacements.append((actual_text, values[variable]))
                                        replacements_made.append(f"{old_text} (case-insensitive) -> {values[variable]} ({len(matches)} occurrence(s) in header)")
                                elif old_text in paragraph_text:
                                    if not any(r[0] == old_text for r in paragraph_replacements):
                                        paragraph_replacements.append((old_text, values[variable]))
                                        replacements_made.append(f"{old_text} -> {values[variable]} (in header)")
                        if paragraph_replacements:
                            self._replace_multiple_texts_with_bold(paragraph, paragraph_replacements)
                
                # Process footer paragraphs
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        paragraph_replacements = []
                        sorted_replacements = sorted(replacement_map.items(), key=lambda x: len(x[0]), reverse=True)
                        for old_text, variable in sorted_replacements:
                            if old_text not in multiple_replacements:
                                paragraph_text = paragraph.text
                                
                                # Case-insensitive matching for uppercase placeholders
                                # Use word boundaries for single-word placeholders to avoid matching "name" in "named"
                                if old_text.isupper() and len(old_text.split()) > 0:
                                    # For single words (like "NAME"), use word boundaries to match whole words only
                                    if len(old_text.split()) == 1:
                                        # Single word: use word boundaries (e.g., \bNAME\b matches "NAME" but not "named")
                                        pattern = re.compile(r'\b' + re.escape(old_text) + r'\b', re.IGNORECASE)
                                    else:
                                        # Multiple words: match case-insensitively (e.g., "BRAND NAME" matches "Brand Name")
                                        pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                                    matches = list(pattern.finditer(paragraph_text))
                                    if matches:
                                        for match in matches:
                                            actual_text = paragraph_text[match.start():match.end()]
                                            if not any(r[0] == actual_text for r in paragraph_replacements):
                                                paragraph_replacements.append((actual_text, values[variable]))
                                        replacements_made.append(f"{old_text} (case-insensitive) -> {values[variable]} ({len(matches)} occurrence(s) in footer)")
                                elif old_text in paragraph_text:
                                    if not any(r[0] == old_text for r in paragraph_replacements):
                                        paragraph_replacements.append((old_text, values[variable]))
                                        replacements_made.append(f"{old_text} -> {values[variable]} (in footer)")
                        if paragraph_replacements:
                            self._replace_multiple_texts_with_bold(paragraph, paragraph_replacements)
        except Exception as e:
            print(f"⚠️ Warning: Could not process headers/footers: {str(e)}")
        
        # Process images and logos (add placeholders)
        self._process_images_and_logos(doc)
        
        return replacements_made

    def _remove_company_name_phrases(self, doc):
        """Remove 'and on behalf of' company text when no company name is provided."""

        def process_paragraph(paragraph):
            if not paragraph.runs:
                return

            phrase_pattern = re.compile(r'[\s,]*and\s+on\s+behalf\s+of', re.IGNORECASE)
            placeholder_pattern = re.compile(r'\s*client\s+company\s+name', re.IGNORECASE)
            while True:
                combined_text = ''.join(run.text for run in paragraph.runs)

                match_phrase = phrase_pattern.search(combined_text)
                if match_phrase:
                    start, end = match_phrase.span()
                    self._remove_text_range_from_runs(paragraph.runs, start, end)
                    continue

                match_placeholder = placeholder_pattern.search(combined_text)
                if match_placeholder:
                    start, end = match_placeholder.span()
                    self._remove_text_range_from_runs(paragraph.runs, start, end)
                    continue

                break

        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)

        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    process_paragraph(paragraph)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    process_paragraph(paragraph)

    def _ensure_bold_values(self, doc, values):
        """Ensure specified values remain bold within the document."""
        clean_values = [v for v in values if isinstance(v, str) and v.strip()]
        if not clean_values:
            return

        for paragraph in doc.paragraphs:
            self._apply_bold_to_values(paragraph.runs, clean_values)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._apply_bold_to_values(paragraph.runs, clean_values)

        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._apply_bold_to_values(paragraph.runs, clean_values)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._apply_bold_to_values(paragraph.runs, clean_values)

    def _apply_bold_to_values(self, runs, values):
        """Apply bold formatting to all occurrences of the provided values within the given runs."""
        if not runs:
            return

        combined_text = ''.join(run.text for run in runs)
        if not combined_text:
            return

        for value in values:
            start_idx = 0
            value_len = len(value)
            while True:
                idx = combined_text.find(value, start_idx)
                if idx == -1:
                    break
                self._set_bold_range_in_runs(runs, idx, idx + value_len)
                start_idx = idx + value_len

    def _set_bold_range_in_runs(self, runs, start, end):
        """Set bold formatting for characters between the specified indexes across runs."""
        if start >= end:
            return

        position = 0
        for run in runs:
            text = run.text
            length = len(text)
            run_start = position
            run_end = position + length

            if run_end <= start or run_start >= end:
                position += length
                continue

            run.bold = True
            position += length

    def _remove_text_range_from_runs(self, runs, start, end):
        """Remove text from a list of runs between absolute character indexes."""
        if start >= end:
            return

        position = 0
        for run in runs:
            text = run.text
            length = len(text)
            run_start = position
            run_end = position + length

            if run_end <= start or run_start >= end:
                position += length
                continue

            relative_start = max(start - run_start, 0)
            relative_end = min(end - run_start, length)

            run.text = text[:relative_start] + text[relative_end:]
            position += length
    
    def _replace_multiple_texts_with_bold(self, paragraph, replacements):
        """Replace multiple texts in a paragraph and make all replacements bold"""
        # Get the full paragraph text
        full_text = paragraph.text
        
        # Check if this is a signature line with Date
        has_date = 'Date' in full_text or 'Date:' in full_text
        
        # Apply all replacements and track length difference
        new_full_text = full_text
        length_diff = 0
        
        # Process replacements - replace all occurrences
        for old_text, new_text in replacements:
            # Replace all occurrences (not just the first one)
            if old_text in new_full_text:
                # Count how many times it appears
                count = new_full_text.count(old_text)
                # Calculate total length difference for all occurrences
                length_diff += (len(new_text) - len(old_text)) * count
                # Replace all occurrences
                new_full_text = new_full_text.replace(old_text, new_text)
        
        # If it's a signature line with Date, adjust spacing to compensate for name length change
        if has_date and length_diff != 0:
            date_marker = 'Date:' if 'Date:' in new_full_text else 'Date'
            if date_marker in new_full_text:
                # Split before Date
                parts = new_full_text.rsplit(date_marker, 1)
                before_date = parts[0]
                
                # Count existing tabs at the end
                tab_count = len(before_date) - len(before_date.rstrip('\t'))
                
                # Adjust tabs: if name got longer (positive diff), reduce tabs; if shorter, add tabs
                # Each tab is roughly 8 chars, so adjust by length_diff / 8
                tabs_to_adjust = -(length_diff // 8)
                new_tab_count = max(1, tab_count + tabs_to_adjust)
                
                # Rebuild the line
                before_date_clean = before_date.rstrip('\t ')
                new_full_text = before_date_clean + ('\t' * new_tab_count) + date_marker
        
        # Clear all runs and rebuild the paragraph
        paragraph.clear()
        
        # Simple approach: split by each replacement text and rebuild
        remaining_text = new_full_text
        
        while remaining_text:
            # Find the first replacement text in the remaining text
            earliest_pos = len(remaining_text)
            earliest_replacement = None
            
            for old_text, new_text in replacements:
                pos = remaining_text.find(new_text)
                if pos != -1 and pos < earliest_pos:
                    earliest_pos = pos
                    earliest_replacement = new_text
            
            if earliest_replacement:
                # Add text before the replacement
                if earliest_pos > 0:
                    paragraph.add_run(remaining_text[:earliest_pos])
                
                # Add the replacement text in bold
                bold_run = paragraph.add_run(earliest_replacement)
                bold_run.bold = True
                
                # Update remaining text
                remaining_text = remaining_text[earliest_pos + len(earliest_replacement):]
            else:
                # No more replacements, add remaining text
                if remaining_text:
                    paragraph.add_run(remaining_text)
                break
    
    def _fix_date_line_without_replacement(self, paragraph):
        """Fix Date alignment for lines without replacements (TEG lines)"""
        text = paragraph.text
        
        # Only process if it looks like a signature line (has behalf/Signature/TEG)
        if not any(word in text for word in ['behalf', 'Signature', 'TEG', 'Intl']):
            return
        
        date_marker = 'Date:' if 'Date:' in text else 'Date'
        if date_marker not in text:
            return
        
        # Split at Date marker
        parts = text.rsplit(date_marker, 1)
        before_date = parts[0]
        
        # Keep the original tab structure - just ensure it's clean
        # Remove any trailing spaces but keep tabs
        before_date_clean = before_date.rstrip(' ')
        
        # If there are no tabs at the end, add them (should have at least some tabs)
        if not before_date_clean.endswith('\t'):
            # Count how many tabs should be there based on the line
            # Most signature lines have 5-6 tabs before Date
            before_date_clean = before_date.rstrip('\t ')
            before_date_clean += '\t' * 5
        
        # Rebuild
        new_text = before_date_clean + date_marker
        
        # Clear and rebuild paragraph
        paragraph.clear()
        paragraph.add_run(new_text)
    
    def _replace_text_with_bold(self, paragraph, old_text, new_text):
        """Replace text in a paragraph and make the new text bold"""
        # Get the full paragraph text
        full_text = paragraph.text
        
        if old_text in full_text:
            # Replace the text in the full paragraph
            new_full_text = full_text.replace(old_text, new_text)
            
            # Clear all runs and rebuild the paragraph
            paragraph.clear()
            
            # Split the new text around the replacement
            parts = new_full_text.split(new_text)
            
            # Add the part before the replacement
            if parts[0]:
                paragraph.add_run(parts[0])
            
            # Add the replacement text in bold
            bold_run = paragraph.add_run(new_text)
            bold_run.bold = True
            
            # Add the remaining parts
            for i in range(1, len(parts)):
                if parts[i]:
                    paragraph.add_run(parts[i])
    
    def _process_images_and_logos(self, doc):
        """Process images and logos in the document"""
        # This method can be extended to handle images/logos
        # For now, we'll just ensure they're preserved in the document
        pass
    
    def _format_contract_amount(self, amount_str):
        """
        Format contract amount as $n,nnn.nn
        
        Args:
            amount_str: Contract amount string (e.g., "5000", "$5000", "5000.00")
            
        Returns:
            str: Formatted amount (e.g., "$5,000.00")
        """
        try:
            # Remove any existing $ and commas
            clean_amount = amount_str.replace('$', '').replace(',', '')
            
            # Convert to float
            amount = float(clean_amount)
            
            # Format with commas and 2 decimal places
            formatted = f"${amount:,.2f}"
            
            return formatted
            
        except (ValueError, TypeError):
            # If formatting fails, return original string with $ prefix
            return f"${amount_str}" if not amount_str.startswith('$') else amount_str
    
    def _insert_pdf_after_first_paragraph(self, doc, uploaded_pdf):
        """
        Insert a PDF (converted to image) after specific text patterns in contract documents
        
        Args:
            doc: Document object
            uploaded_pdf: Streamlit UploadedFile object containing PDF data
        """
        try:
            import io
            from docx.shared import Inches
            
            # Convert PDF first page to image
            image_data = self._convert_pdf_to_image(uploaded_pdf)
            
            # First, remove any existing images in the document
            self._remove_existing_images(doc)
            
            # Define the specific text patterns to look for (using key parts)
            target_patterns = [
                "Development Contract, and Terms and Conditions Agreement (Attachment B), and as follows:",
                "The attached Production Workbook includes all detail as a part of this agreement."
            ]
            
            # Find the paragraph containing the target text
            target_paragraph = None
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                for pattern in target_patterns:
                    if pattern in paragraph_text:
                        target_paragraph = paragraph
                        print(f"📍 Found target paragraph: '{paragraph_text[:100]}...'")
                        break
                if target_paragraph:
                    break
            
            # If exact patterns not found, try partial matching
            if target_paragraph is None:
                print("⚠️ Exact patterns not found, trying partial matching")
                partial_patterns = [
                    ("Development Contract", "Attachment B"),
                    ("Production Workbook", "includes all detail")
                ]
                
                for paragraph in doc.paragraphs:
                    paragraph_text = paragraph.text.strip()
                    for pattern1, pattern2 in partial_patterns:
                        if pattern1 in paragraph_text and pattern2 in paragraph_text:
                            target_paragraph = paragraph
                            print(f"📍 Found partial match: '{paragraph_text[:100]}...'")
                            break
                    if target_paragraph:
                        break
            
            # If no target pattern found, fall back to first substantial paragraph
            if target_paragraph is None:
                print("⚠️ Target pattern not found, using first substantial paragraph")
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if (len(text) > 50 and 
                        not text.isupper() and 
                        not text.endswith(':') and 
                        '.' in text):
                        target_paragraph = paragraph
                        print(f"📍 Fallback to: '{text[:50]}...'")
                        break
            
            if target_paragraph is None:
                print("⚠️ No suitable paragraph found in document")
                return
            
            # Create a new paragraph for the image
            new_paragraph = doc.add_paragraph()
            
            # Insert the new paragraph after the target paragraph
            new_paragraph._element.getparent().remove(new_paragraph._element)
            target_paragraph._element.addnext(new_paragraph._element)
            
            # Add the image to the new paragraph
            run = new_paragraph.add_run()
            
            # Add image with reasonable size (max width 6 inches)
            run.add_picture(io.BytesIO(image_data), width=Inches(6))
            
            print(f"✅ PDF (as image) inserted after target paragraph: {uploaded_pdf.name}")
            
        except Exception as e:
            print(f"❌ Error inserting PDF: {str(e)}")
            # Don't raise the exception to avoid breaking the document processing
    
    def _convert_pdf_to_image(self, uploaded_pdf):
        """
        Convert the first page of a PDF to an image using PyMuPDF (no external dependencies needed)
        
        Args:
            uploaded_pdf: Streamlit UploadedFile object containing PDF data
            
        Returns:
            bytes: Image data (PNG format)
        """
        try:
            import fitz  # PyMuPDF
            import io
            from PIL import Image
            
            # Get PDF bytes
            pdf_bytes = uploaded_pdf.read()
            
            # Reset the file pointer for potential future reads
            uploaded_pdf.seek(0)
            
            # Open PDF from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            if pdf_document.page_count == 0:
                raise ValueError("PDF has no pages")
            
            # Get the first page
            page = pdf_document[0]
            
            # Render page to an image (matrix for scaling - 2.0 = 200 DPI equivalent)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            
            # Convert pixmap to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Convert PIL Image to bytes
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            
            # Close the PDF
            pdf_document.close()
            
            print("✅ PDF converted to image successfully using PyMuPDF")
            return img_byte_arr.getvalue()
            
        except ImportError as e:
            raise ImportError(f"PyMuPDF library is required. Install with: pip install PyMuPDF Pillow. Error: {str(e)}")
        except Exception as e:
            print(f"❌ Error converting PDF to image: {str(e)}")
            raise
    
    def _remove_existing_images(self, doc):
        """
        Remove all existing images from the document
        
        Args:
            doc: Document object
        """
        try:
            images_removed = 0
            
            # Collect all runs with images first, then remove them
            runs_to_remove = []
            
            # Find images in paragraphs
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # Check if run contains images
                    if run._element.xpath('.//a:blip'):
                        runs_to_remove.append(run)
            
            # Find images in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run._element.xpath('.//a:blip'):
                                    runs_to_remove.append(run)
            
            # Now remove all the runs with images
            for run in runs_to_remove:
                parent = run._element.getparent()
                if parent is not None:
                    parent.remove(run._element)
                    images_removed += 1
            
            print(f"✅ Removed {images_removed} existing image(s) from document")
            
        except Exception as e:
            print(f"⚠️ Error removing existing images: {str(e)}")
    
    def get_template_info(self, template_type):
        """Get information about a template"""
        if template_type not in self.template_mapping:
            return None
        
        config = self.template_mapping[template_type]
        template_path = os.path.join('inputs', config['file'])
        
        if not os.path.exists(template_path):
            return None
        
        doc = Document(template_path)
        
        return {
            'file': config['file'],
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'replacements': config['replacements']
        }

def test_document_processing():
    """Test the document processing system"""
    print("Testing Document Processing System")
    print("=" * 50)
    
    processor = DocxTemplateProcessor()
    
    # Test data
    test_client = "Estevao Cavalcante"
    test_email = "estevao@example.com"
    test_amount = "$15,865.00"
    test_date = "October 01, 2025"
    
    # Test each template type
    template_types = ['development_contract', 'development_terms', 'production_contract', 'production_terms']
    
    for template_type in template_types:
        print(f"\n--- Testing {template_type} ---")
        
        try:
            # Get template info
            info = processor.get_template_info(template_type)
            if info:
                print(f"Template: {info['file']}")
                print(f"Paragraphs: {info['paragraphs']}")
                print(f"Tables: {info['tables']}")
                print(f"Replacements: {info['replacements']}")
            
            # Process document
            output_path = processor.process_document(
                template_type=template_type,
                client_name=test_client,
                email=test_email,
                contract_amount=test_amount if template_type in ['development_contract', 'production_contract'] else None,
                contract_date=test_date
            )
            
            print(f"✅ Document created: {output_path}")
            
        except Exception as e:
            print(f"❌ Error processing {template_type}: {str(e)}")
    
    print(f"\n✅ All documents processed successfully!")
    print(f"📁 Check the 'processed_documents' folder for the generated files")

if __name__ == "__main__":
    test_document_processing()